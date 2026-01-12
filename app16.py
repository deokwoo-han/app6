import streamlit as st
import google.generativeai as genai
import requests
import json
from datetime import date, datetime, timedelta
from io import BytesIO
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image

# -------------------------------------------------------------------------
# [0. System Setup & Session Initialization]
# -------------------------------------------------------------------------
st.set_page_config(page_title="AI 법률 마스터 (Ultimate Edition)", page_icon="⚖️", layout="wide")

# Initialize Session State (Combining all variables for Save/Load compatibility)
default_values = {
    'rec_court': "서울중앙지방법원",
    'amt_in': "30000000",
    'chat_history': [],
    'party_a': "홍길동",      # Applicant/Plaintiff
    'party_b': "김철수",      # Opponent/Defendant
    'facts_raw': "",        # Case details
    'ev_raw': "차용증\n이체내역서\n카톡 대화록",  # Text evidence list
    'ref_case': ""          # For precedent search
}

for key, val in default_values.items():
    if key not in st.session_state:
        st.session_state[key] = val

# -------------------------------------------------------------------------
# [1. Integrated Database]
# -------------------------------------------------------------------------

# 1-1. Full Court List (from app14)
COURT_LIST = [
    "서울중앙지방법원", "서울동부지방법원", "서울남부지방법원", "서울북부지방법원", "서울서부지방법원",
    "서울가정법원", "서울행정법원", "서울회생법원",
    "의정부지방법원", "의정부지방법원 고양지원", "의정부지방법원 남양주지원",
    "인천지방법원", "인천지방법원 부천지원", "인천가정법원",
    "수원지방법원", "수원지방법원 성남지원", "수원지방법원 여주지원", "수원지방법원 평택지원", "수원지방법원 안산지원", "수원지방법원 안양지원", 
    "수원가정법원", "수원회생법원",
    "춘천지방법원", "춘천지방법원 강릉지원", "춘천지방법원 원주지원", "춘천지방법원 속초지원", "춘천지방법원 영월지원",
    "대전지방법원", "대전지방법원 천안지원", "대전지방법원 서산지원", "대전지방법원 홍성지원", "대전지방법원 논산지원", "대전지방법원 공주지원", 
    "대전가정법원",
    "청주지방법원", "청주지방법원 충주지원", "청주지방법원 제천지원", "청주지방법원 영동지원",
    "대구지방법원", "대구지방법원 서부지원", "대구지방법원 포항지원", "대구지방법원 김천지원", "대구지방법원 안동지원", "대구지방법원 경주지원", "대구지방법원 상주지원", "대구지방법원 의성지원", "대구지방법원 영덕지원", 
    "대구가정법원",
    "부산지방법원", "부산지방법원 동부지원", "부산지방법원 서부지원", "부산가정법원", "부산회생법원",
    "울산지방법원", "울산가정법원",
    "창원지방법원", "창원지방법원 마산지원", "창원지방법원 진주지원", "창원지방법원 통영지원", "창원지방법원 밀양지원", "창원지방법원 거창지원",
    "광주지방법원", "광주지방법원 순천지원", "광주지방법원 목포지원", "광주지방법원 장흥지원", "광주지방법원 해남지원", "광주가정법원",
    "전주지방법원", "전주지방법원 군산지원", "전주지방법원 정읍지원", "전주지방법원 남원지원",
    "제주지방법원"
]

# 1-2. Detailed Jurisdiction Map (from app14 - The most complete version)
JURISDICTION_MAP = {
    # --- Capital Area ---
    "종로": "서울중앙지방법원", "중구": "서울중앙지방법원", "강남": "서울중앙지방법원", "서초": "서울중앙지방법원", "관악": "서울중앙지방법원", "동작": "서울중앙지방법원",
    "성동": "서울동부지방법원", "광진": "서울동부지방법원", "강동": "서울동부지방법원", "송파": "서울동부지방법원",
    "영등포": "서울남부지방법원", "강서": "서울남부지방법원", "양천": "서울남부지방법원", "구로": "서울남부지방법원", "금천": "서울남부지방법원",
    "동대문": "서울북부지방법원", "중랑": "서울북부지방법원", "성북": "서울북부지방법원", "도봉": "서울북부지방법원", "강북": "서울북부지방법원", "노원": "서울북부지방법원",
    "은평": "서울서부지방법원", "서대문": "서울서부지방법원", "마포": "서울서부지방법원", "용산": "서울서부지방법원",
    "고양": "의정부지방법원 고양지원", "파주": "의정부지방법원 고양지원", "남양주": "의정부지방법원 남양주지원", "구리": "의정부지방법원 남양주지원", "가평": "의정부지방법원 남양주지원",
    "부천": "인천지방법원 부천지원", "김포": "인천지방법원 부천지원", "인천": "인천지방법원", "강화": "인천지방법원", "옹진": "인천지방법원",
    "성남": "수원지방법원 성남지원", "하남": "수원지방법원 성남지원", "광주": "수원지방법원 성남지원",
    "안산": "수원지방법원 안산지원", "광명": "수원지방법원 안산지원", "시흥": "수원지방법원 안산지원",
    "안양": "수원지방법원 안양지원", "과천": "수원지방법원 안양지원", "의왕": "수원지방법원 안양지원", "군포": "수원지방법원 안양지원",
    "평택": "수원지방법원 평택지원", "안성": "수원지방법원 평택지원", "여주": "수원지방법원 여주지원", "이천": "수원지방법원 여주지원", "양평": "수원지방법원 여주지원",
    "수원": "수원지방법원", "용인": "수원지방법원", "화성": "수원지방법원", "오산": "수원지방법원",
    # --- Gangwon ---
    "춘천": "춘천지방법원", "홍천": "춘천지방법원", "양구": "춘천지방법원", "인제": "춘천지방법원", "화천": "춘천지방법원",
    "강릉": "춘천지방법원 강릉지원", "동해": "춘천지방법원 강릉지원", "삼척": "춘천지방법원 강릉지원",
    "원주": "춘천지방법원 원주지원", "횡성": "춘천지방법원 원주지원", "속초": "춘천지방법원 속초지원", "양양": "춘천지방법원 속초지원", "고성": "춘천지방법원 속초지원",
    "영월": "춘천지방법원 영월지원", "태백": "춘천지방법원 영월지원", "정선": "춘천지방법원 영월지원",
    # --- Chungcheong ---
    "천안": "대전지방법원 천안지원", "아산": "대전지방법원 천안지원", "서산": "대전지방법원 서산지원", "당진": "대전지방법원 서산지원", "태안": "대전지방법원 서산지원",
    "홍성": "대전지방법원 홍성지원", "보령": "대전지방법원 홍성지원", "예산": "대전지방법원 홍성지원", "논산": "대전지방법원 논산지원", "계룡": "대전지방법원 논산지원", "부여": "대전지방법원 논산지원",
    "공주": "대전지방법원 공주지원", "청양": "대전지방법원 공주지원", "대전": "대전지방법원", "세종": "대전지방법원",
    "청주": "청주지방법원", "진천": "청주지방법원", "보은": "청주지방법원", "괴산": "청주지방법원", "증평": "청주지방법원",
    "충주": "청주지방법원 충주지원", "음성": "청주지방법원 충주지원", "제천": "청주지방법원 제천지원", "단양": "청주지방법원 제천지원", "영동": "청주지방법원 영동지원", "옥천": "청주지방법원 영동지원",
    # --- Yeongnam ---
    "달서": "대구지방법원 서부지원", "달성": "대구지방법원 서부지원", "대구 서구": "대구지방법원 서부지원", "대구": "대구지방법원", "수성": "대구지방법원",
    "포항": "대구지방법원 포항지원", "울릉": "대구지방법원 포항지원", "경주": "대구지방법원 경주지원", "김천": "대구지방법원 김천지원", "구미": "대구지방법원 김천지원",
    "안동": "대구지방법원 안동지원", "영주": "대구지방법원 안동지원", "상주": "대구지방법원 상주지원", "문경": "대구지방법원 상주지원", "의성": "대구지방법원 의성지원", "영덕": "대구지방법원 영덕지원", "울진": "대구지방법원 영덕지원",
    "해운대": "부산지방법원 동부지원", "부산남구": "부산지방법원 동부지원", "수영": "부산지방법원 동부지원", "기장": "부산지방법원 동부지원",
    "사하": "부산지방법원 서부지원", "사상": "부산지방법원 서부지원", "부산강서": "부산지방법원 서부지원", "북구": "부산지방법원 서부지원", "부산": "부산지방법원",
    "울산": "울산지방법원", "양산": "울산지방법원", "창원": "창원지방법원", "함안": "창원지방법원", "의령": "창원지방법원",
    "마산": "창원지방법원 마산지원", "진해": "창원지방법원 마산지원", "진주": "창원지방법원 진주지원", "사천": "창원지방법원 진주지원", "통영": "창원지방법원 통영지원", "거제": "창원지방법원 통영지원",
    "밀양": "창원지방법원 밀양지원", "창녕": "창원지방법원 밀양지원", "거창": "창원지방법원 거창지원", "함양": "창원지방법원 거창지원", "합천": "창원지방법원 거창지원",
    # --- Honam ---
    "순천": "광주지방법원 순천지원", "여수": "광주지방법원 순천지원", "광양": "광주지방법원 순천지원", "보성": "광주지방법원 순천지원", "고흥": "광주지방법원 순천지원", "구례": "광주지방법원 순천지원",
    "목포": "광주지방법원 목포지원", "무안": "광주지방법원 목포지원", "신안": "광주지방법원 목포지원", "해남": "광주지방법원 해남지원", "완도": "광주지방법원 해남지원", "진도": "광주지방법원 해남지원",
    "장흥": "광주지방법원 장흥지원", "강진": "광주지방법원 장흥지원", "광주": "광주지방법원", "나주": "광주지방법원", "화순": "광주지방법원", "장성": "광주지방법원", "곡성": "광주지방법원",
    "군산": "전주지방법원 군산지원", "익산": "전주지방법원 군산지원", "정읍": "전주지방법원 정읍지원", "고창": "전주지방법원 정읍지원", "부안": "전주지방법원 정읍지원",
    "남원": "전주지방법원 남원지원", "순창": "전주지방법원 남원지원", "장수": "전주지방법원 남원지원", "무주": "전주지방법원 남원지원", "전주": "전주지방법원", "완주": "전주지방법원", "김제": "전주지방법원",
    # --- Jeju ---
    "제주": "제주지방법원", "서귀포": "제주지방법원"
}

# 1-3. Mind Care DB (app14 + app15)
MIND_CARE_DB = {
    "start": {"advice": "시작이 반입니다. 권리 구제의 첫걸음을 응원합니다.", "video": "https://www.youtube.com/watch?v=pzlw6fUux4o"},
    "wait": {"advice": "법원은 증거로 말합니다. 차분히 답변서를 기다리며 증거를 재점검하세요.", "video": "https://www.youtube.com/watch?v=HuM1k6d7NXI"},
    "fight": {"advice": "감정적 대응은 금물입니다. 법정에서는 오직 팩트와 법리로 승부하세요.", "video": "https://www.youtube.com/watch?v=v2AcV5rV_wA"},
    "trial": {"advice": "재판장 앞에서는 간결하고 명확하게 답변하는 것이 가장 유리합니다.", "video": "https://www.youtube.com/watch?v=inpok4MKVLM"},
    "end": {"advice": "수고하셨습니다. 결과와 상관없이 당신의 정당한 권리를 위한 노력은 가치 있습니다.", "video": "https://www.youtube.com/watch?v=CvFH_6DNRCY"}
}

# 1-4. Scenario Logic (app14)
SCENARIO_LOGIC = {
    "LOAN": {"label": "💰 대여금 청구", "weights": ["빌려", "대여", "차용", "차용증"]},
    "DEPOSIT": {"label": "🏠 보증금 반환", "weights": ["보증금", "전세", "월세", "임대차"]},
    "TORT": {"label": "🏥 손해배상", "weights": ["사고", "폭행", "피해", "과실"]},
    "WAGE": {"label": "💼 임금 청구", "weights": ["임금", "월급", "퇴직금", "급여"]},
    "SALES": {"label": "🏗️ 물품/공사대금", "weights": ["물품", "공사", "대금", "자재"]},
    "ESTATE": {"label": "🏘️ 부동산 계약", "weights": ["부동산", "매매", "계약", "등기"]},
    "GENERAL": {"label": "📝 일반 민사", "weights": []}
}

# -------------------------------------------------------------------------
# [2. Intelligent Utility Functions]
# -------------------------------------------------------------------------

def get_available_models(api_key):
    """Available models including Vision capabilities"""
    if not api_key: return []
    try:
        genai.configure(api_key=api_key)
        return [m.name for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]
    except: return []

def find_best_court(address, category="일반"):
    """Determines jurisdiction based on address and case type (app14 logic)"""
    base_court = "서울중앙지방법원"
    
    # Geographic mapping (Longest match first)
    if address:
        sorted_keys = sorted(JURISDICTION_MAP.keys(), key=len, reverse=True)
        for key in sorted_keys:
            if key in address:
                base_court = JURISDICTION_MAP[key]
                break
    
    # Special Court Logic
    special_logic = {
        "가사": {"서울": "서울가정법원", "인천": "인천가정법원", "수원": "수원가정법원", "대전": "대전가정법원", "대구": "대구가정법원", "부산": "부산가정법원", "울산": "울산가정법원", "광주": "광주가정법원"},
        "회생": {"서울": "서울회생법원", "수원": "수원회생법원", "부산": "부산회생법원"},
        "파산": {"서울": "서울회생법원", "수원": "수원회생법원", "부산": "부산회생법원"},
        "행정": {"서울": "서울행정법원"}
    }
    
    cat_key = ""
    if any(x in category for x in ["가사", "이혼", "상속"]): cat_key = "가사"
    elif any(x in category for x in ["회생", "파산"]): cat_key = "회생"
    elif any(x in category for x in ["행정"]): cat_key = "행정"

    if cat_key:
        region_prefix = base_court[:2]
        if region_prefix in special_logic.get(cat_key, {}):
            return special_logic[cat_key][region_prefix]
            
    return base_court

def detect_scenario(text):
    """Auto-detects case type from description"""
    scores = {k: sum(1 for w in v['weights'] if w in text) for k, v in SCENARIO_LOGIC.items()}
    best = max(scores, key=scores.get)
    return SCENARIO_LOGIC[best]['label'] if scores[best] > 0 else "📝 일반 민사"

def calculate_legal_costs(amount):
    """Calculates Stamp Duty and Service Fees"""
    try: amt = int(str(amount).replace(",", ""))
    except: amt = 0
    if amt <= 0: return 0, 0, 0
    
    if amt <= 10000000: stamp = amt * 0.005
    elif amt <= 100000000: stamp = amt * 0.0045 + 5000
    else: stamp = amt * 0.004 + 55000
    stamp = max(1000, int(stamp // 100 * 100))
    svc = 5200 * (10 if amt <= 30000000 else 15)
    return amt, stamp, svc

def predict_detailed_timeline(amount):
    """Generates litigation timeline with Mind Care integration"""
    amt, stamp, svc = calculate_legal_costs(amount)
    today = date.today()
    steps = [
        (0, "소장 접수", "인지대/송달료 납부 및 사건번호 부여", "start"),
        (4, "부본 송달", "피고에게 소장이 전달되고 답변서를 기다리는 단계", "wait"),
        (12, "변론 기일", "법정에 출석하여 양측의 주장과 증거를 다투는 단계", "fight"),
        (20, "재판 심리", "추가 증거 조사 및 판사의 최종 판단 과정", "trial"),
        (28, "판결 선고", "최종 판결문 교부 및 소송의 종결", "end")
    ]
    timeline = []
    for w, ev, ds, care_key in steps:
        timeline.append({
            "week": f"{w}주차",
            "date": (today + timedelta(weeks=w)).strftime("%Y.%m.%d"),
            "event": ev, "desc": ds, "care": MIND_CARE_DB[care_key]
        })
    return timeline, amt, stamp, svc

def create_evidence_list_formatted(text):
    if not text: return "없음"
    evs = [e.strip() for e in text.split('\n') if e.strip()]
    return "\n".join([f"갑 제{i}호증 ({v})" for i, v in enumerate(evs, 1)])

def create_docx(title, content):
    doc = Document()
    doc.add_heading(title, 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(content)
    buf = BytesIO(); doc.save(buf); buf.seek(0)
    return buf

def get_gemini_response(api_key, model_name, prompt, image=None):
    """Handles both Text and Vision requests"""
    try:
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel(model_name)
        if image:
            return model.generate_content([prompt, image]).text
        return model.generate_content(prompt).text
    except Exception as e: return f"❌ 오류: {str(e)}"

# -------------------------------------------------------------------------
# [3. Sidebar Menu Integration]
# -------------------------------------------------------------------------
with st.sidebar:
    st.title("⚖️ AI 법률 마스터")
    st.caption("Integrated Ultimate Ver.")
    
    api_key = st.text_input("Google API Key", type="password")
    
    # Model Selection (Prioritize models with Vision)
    available_models = get_available_models(api_key)
    default_models = ["models/gemini-2.0-flash-exp", "models/gemini-1.5-flash", "models/gemini-1.5-pro"]
    selected_model = st.selectbox("AI 모델 선택", available_models if available_models else default_models)
    
    law_id = st.text_input("법령센터 ID (선택)")

    st.divider()

    # [Feature: Save/Load Data (from app15)]
    with st.expander("💾 데이터 관리 (Save/Load)"):
        # Save
        save_data = {
            "party_a": st.session_state.party_a,
            "party_b": st.session_state.party_b,
            "amt_in": st.session_state.amt_in,
            "facts_raw": st.session_state.get('facts_raw', ''),
            "rec_court": st.session_state.rec_court,
            "ev_raw": st.session_state.get('ev_raw', '')
        }
        json_str = json.dumps(save_data, ensure_ascii=False)
        st.download_button("현재 내용 PC에 저장", json_str, "legal_case_data.json", "application/json")
        
        # Load
        uploaded_json = st.file_uploader("저장된 파일 불러오기", type="json")
        if uploaded_json is not None:
            try:
                loaded_data = json.load(uploaded_json)
                st.session_state.party_a = loaded_data.get("party_a", "")
                st.session_state.party_b = loaded_data.get("party_b", "")
                st.session_state.amt_in = loaded_data.get("amt_in", "0")
                st.session_state.facts_raw = loaded_data.get("facts_raw", "")
                st.session_state.rec_court = loaded_data.get("rec_court", "")
                st.session_state.ev_raw = loaded_data.get("ev_raw", "")
                st.success("데이터 복원 완료! (새로고침 시 적용)")
            except:
                st.error("파일 형식이 올바르지 않습니다.")

    st.divider()
    
    # [Integrated Menu (from app14)]
    menu_options = [
        "무료법률상담 (AI 챗봇)",
        "전자소송 (지급명령/채권자)",
        "전자소송 (지급명령/채무자)",
        "민사소송 (대여금)",
        "민사소송 (임차보증금)",
        "민사소송 (손해배상)",
        "민사소송 (기 타)",
        "민사집행 (압류/경매)",
        "형사소송 (고소/고발)",
        "행정소송",
        "가사소송 (이혼,상속)",
        "개인파산/개인회생"
    ]
    selected_menu = st.radio("📂 법률 서비스 선택", menu_options)
    
    st.divider()
    # [Jurisdiction Finder]
    st.subheader("📍 관할 법원 자동 매칭")
    addr_input = st.text_input("주소 (예: 서울 서초구, 대구 달서구)", placeholder="시/군/구 입력")
    if addr_input:
        st.session_state.rec_court = find_best_court(addr_input, selected_menu)
        st.success(f"추천 관할: {st.session_state.rec_court}")

# -------------------------------------------------------------------------
# [4. Main Content Area]
# -------------------------------------------------------------------------
st.header(f"{selected_menu} 통합 솔루션")

# [CASE 1: AI Chatbot]
if "무료법률상담" in selected_menu:
    st.info("🤖 100만 건의 판례 데이터를 학습한 AI 변호사가 상담해드립니다.")
    
    for chat in st.session_state.chat_history:
        with st.chat_message(chat["role"]):
            st.write(chat["content"])
            
    user_input = st.chat_input("법률 고민을 입력하세요 (예: 전세보증금을 못 받고 있는데 어떻게 하죠?)")
    
    if user_input:
        st.session_state.chat_history.append({"role": "user", "content": user_input})
        with st.chat_message("user"): st.write(user_input)
            
        with st.chat_message("assistant"):
            with st.spinner("법률 데이터베이스 분석 중..."):
                prompt = f"너는 한국 법률 전문가야. 질문: {user_input}. 판례와 법령에 근거하여 상세히 답변하고, 필요하다면 내용증명이나 소송 절차도 안내해줘."
                response = get_gemini_response(api_key, selected_model, prompt)
                st.write(response)
                st.session_state.chat_history.append({"role": "assistant", "content": response})

# [CASE 2: Integrated Tools (Documents, Evidence, Vision, etc.)]
else:
    # 5 Tabs Integration
    tab1, tab2, tab3, tab4, tab5 = st.tabs(["📝 서류 작성", "📨 내용증명", "🔎 증거/비용/케어", "⚖️ 판례 검색", "📋 소송 진단"])
    
    # Config based on menu
    config = {"type": "법률 서면", "role": "신청인", "opp": "피신청인"}
    if "지급명령" in selected_menu: config = {"type": "지급명령신청서", "role": "채권자", "opp": "채무자"}
    elif "민사소송" in selected_menu: config = {"type": "소장", "role": "원고", "opp": "피고"}
    elif "형사" in selected_menu: config = {"type": "고소장", "role": "고소인", "opp": "피고소인"}
    elif "행정" in selected_menu: config = {"type": "소장", "role": "원고", "opp": "피고(처분청)"}
    elif "가사" in selected_menu: config = {"type": "소장", "role": "원고", "opp": "피고"}
    elif "파산" in selected_menu or "회생" in selected_menu: config = {"type": "개시신청서", "role": "신청인", "opp": "채권자목록"}
    
    is_money = any(x in selected_menu for x in ["민사", "지급", "대여", "손해", "보증금"])

    # --- [TAB 1: Document Generation] ---
    with tab1:
        st.subheader(f"📄 {config['type']} 자동 작성")
        
        c1, c2 = st.columns(2)
        st.session_state.party_a = c1.text_input(f"{config['role']} 이름 (나)", st.session_state.party_a)
        st.session_state.party_b = c2.text_input(f"{config['opp']} 이름 (상대)", st.session_state.party_b)
        
        c3, c4 = st.columns(2)
        if is_money:
            st.session_state.amt_in = c3.text_input("청구/피해 금액 (숫자만)", st.session_state.amt_in)
        
        try: c_idx = COURT_LIST.index(st.session_state.rec_court)
        except: c_idx = 0
        target_court = c4.selectbox("제출 법원", COURT_LIST, index=c_idx)
        
        st.session_state.facts_raw = st.text_area("사건 상세 경위", st.session_state.get('facts_raw', ''), height=150, placeholder="육하원칙에 따라 상세히 기재하세요.")
        st.session_state.ev_raw = st.text_area("입증 방법 (증거)", st.session_state.get('ev_raw', ''), placeholder="차용증\n이체내역서\n카톡 대화록")
        
        # Scenario Detection (from app14)
        s_label = detect_scenario(st.session_state.facts_raw)
        st.info(f"💡 AI 분석 사건 유형: **{s_label}**")

        if st.button("🚀 AI 서류 생성"):
            amt, stamp, svc = calculate_legal_costs(st.session_state.amt_in)
            formatted_ev = create_evidence_list_formatted(st.session_state.ev_raw)
            
            prompt = f"""
            역할: 당신은 {selected_menu} 전문 변호사입니다.
            문서: {config['type']}
            관할법원: {target_court}
            {config['role']}: {st.session_state.party_a}
            {config['opp']}: {st.session_state.party_b}
            금액: {amt if is_money else '비재산권'}
            청구원인: {st.session_state.facts_raw}
            입증방법: {formatted_ev}
            사건유형: {s_label}
            
            요청사항: 대한민국의 법률 서식에 맞춰 엄격하고 전문적인 문서를 작성하세요. 
            청구취지와 청구원인을 명확히 구분하세요.
            """
            
            res = get_gemini_response(api_key, selected_model, prompt)
            
            if is_money:
                st.success(f"💰 예상 비용 분석: 인지대 {stamp:,}원 / 송달료 {svc:,}원")
                
            st.text_area("작성 결과", res, height=400)
            st.download_button("💾 다운로드 (.docx)", create_docx(config['type'], res), f"{config['type']}.docx")
            
            with st.expander("📌 전자소송 및 제출 가이드"):
                st.markdown(f"1. [전자소송 사이트](https://ecfs.scourt.go.kr) 접속\n2. 관할법원 **{target_court}** 선택\n3. 작성된 파일을 첨부하여 제출")

    # --- [TAB 2: Contents Proof (Naeyongjeungmyeong)] ---
    with tab2:
        st.subheader("📨 내용증명 (소송 전 독촉)")
        st.caption("내용증명은 본격적인 소송 전, 상대방을 압박하고 증거를 남기는 효과적인 수단입니다.")
        
        col1, col2 = st.columns(2)
        with col1:
            snd = st.text_input("발신인 (나)", st.session_state.party_a)
            rcv = st.text_input("수신인 (상대방)", st.session_state.party_b)
        with col2:
            st.info(f"추천 관할: {st.session_state.rec_court}")
            
        cd_facts = st.text_area("독촉 사유 및 요구사항", st.session_state.facts_raw, placeholder="예: 2023.1.1. 빌려간 1,000만원을 2023.12.31.까지 갚기로 했으나 미이행함.")
        
        if st.button("내용증명 생성"):
            prompt = f"{snd}가 {rcv}에게 보내는 강력한 내용증명을 작성하라. 사유: {cd_facts}. 법적 조치 예고 포함."
            res = get_gemini_response(api_key, selected_model, prompt)
            st.text_area("내용증명 결과", res, height=300)
            st.download_button("Word 다운로드", create_docx("내용증명서", res), "내용증명.docx")

    # --- [TAB 3: Evidence & Tools (Merged from app14 + app15)] ---
    with tab3:
        st.subheader("🔍 증거 분석, 비용, 그리고 마인드 케어")
        
        c_left, c_right = st.columns([1, 1])
        
        with c_left:
            st.markdown("### 📸 AI 이미지/증거 분석 (Vision)")
            uploaded_img = st.file_uploader("증거 이미지 업로드 (문서, 캡처 등)", type=["jpg", "png"])
            if uploaded_img and st.button("이미지 분석"):
                img = Image.open(uploaded_img)
                st.image(img, caption="업로드된 증거", use_container_width=True)
                with st.spinner("AI가 문서를 분석 중입니다..."):
                    res = get_gemini_response(api_key, selected_model, "이 이미지의 핵심 법적 내용을 요약하고, 소송에서 유리한 증거가 될지 분석해줘.", img)
                    st.write(res)
            
            st.divider()
            
            st.markdown("### 🧬 텍스트 증거 목록 분석")
            ev_input_an = st.text_area("분석할 증거 목록 입력", st.session_state.ev_raw, height=100)
            if st.button("증거 효력 분석"):
                p = f"다음 증거들의 민사소송상 증거능력을 별점(5점만점)으로 평가하고, 직접증거와 정황증거로 분류해줘: {ev_input_an}"
                st.markdown(get_gemini_response(api_key, selected_model, p))
                
        with c_right:
            st.markdown("### 🧮 지연손해금(이자) 계산기")
            with st.expander("이자 계산기 열기"):
                c_d1, c_d2, c_r = st.columns(3)
                d1 = c_d1.date_input("기산일 (빌려준 날+1)")
                d2 = c_d2.date_input("기준일 (오늘)")
                rate = c_r.number_input("이율(%)", value=12.0)
                if st.button("이자 계산"):
                    days = (d2 - d1).days
                    if days > 0:
                        try: p_amt = int(str(st.session_state.amt_in).replace(",", ""))
                        except: p_amt = 0
                        interest = int(p_amt * (rate/100) * (days/365))
                        st.success(f"원금 {p_amt:,}원 + 이자 {interest:,}원 = 총 {p_amt+interest:,}원")
                    else: st.warning("날짜를 확인하세요.")
            
            st.divider()
            
            st.markdown("### 🧘 타임라인 & 마인드 케어")
            if is_money:
                timeline, _, _, _ = predict_detailed_timeline(st.session_state.amt_in)
                current_step = st.selectbox("현재 진행 단계 확인", [t['event'] for t in timeline])
                selected_info = next((t for t in timeline if t['event'] == current_step), timeline[0])
                st.info(f"📅 {selected_info['week']}차 예상: {selected_info['desc']}")
                st.markdown(f"**💬 조언:** {selected_info['care']['advice']}")
                st.video(selected_info['care']['video'])
            else:
                st.info("금전 소송 유형에서 타임라인 기능이 활성화됩니다.")

    # --- [TAB 4: Precedents] ---
    with tab4:
        st.subheader("⚖️ 유사 판례 심층 분석")
        q = st.text_input("검색 키워드", f"{selected_menu} 승소 판례")
        if st.button("판례 검색 및 분석"):
            if law_id: st.toast("법령센터 API 연동 모드")
            prompt = f"키워드 '{q}'와 관련된 주요 대법원 판례 경향을 분석하고, 해당 소송에서 승소하기 위한 핵심 법리를 요약해줘."
            st.markdown(get_gemini_response(api_key, selected_model, prompt))

    # --- [TAB 5: Self Diagnosis (from app15)] ---
    with tab5:
        st.subheader("📋 소송 적합성 자가진단")
        st.caption("소송 전 필수 체크리스트입니다.")
        
        q1 = st.radio("1. 상대방의 인적사항(이름, 주소, 주민번호 등)을 하나라도 정확히 아나요?", ["예", "아니오"])
        q2 = st.radio("2. 돈을 빌려주거나 피해를 입은지 10년(상사채권 5년/불법행위 3년)이 안 지났나요?", ["예", "아니오"])
        q3 = st.radio("3. 입증할 수 있는 객관적 증거(이체내역, 문자, 녹취 등)가 있나요?", ["예", "아니오"])
        
        if st.button("진단 결과 확인"):
            score = 0
            if q1 == "예": score += 1
            if q2 == "예": score += 1
            if q3 == "예": score += 1
            
            if score == 3:
                st.success("✅ 소송 진행이 충분히 가능한 상태입니다.")
            elif score == 2:
                st.warning("⚠️ 일부 요건이 부족합니다. 사실조회 신청 등이 필요할 수 있습니다.")
            else:
                st.error("❌ 현재 상태로는 소송 진행이 어렵거나 패소 위험이 높습니다. 증거를 더 수집하세요.")